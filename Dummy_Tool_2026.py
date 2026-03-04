import tkinter as tk
from tkinter import filedialog, messagebox, StringVar
from docx import Document
import xml.etree.ElementTree as ET
from xml.etree.ElementTree import SubElement
import re
import langdetect
import datefinder
import os
import pandas as pd
from tkinter import ttk

def find_dates(text):
    dates = list(datefinder.find_dates(text))
    return dates

def detect_language_from_paragraphs(paragraphs, start, end):
    try:
        combined_text = " ".join(paragraph.text for paragraph in paragraphs[start:end])
        lang = langdetect.detect(combined_text)
        return lang
    except:
        return None

def replace_non_breakable_space(title):
    pattern = r'(\d{1,2})(\s+)?(janvier|février|mars|avril|peut|juin|juillet|août|septembre|octobre|novembre|décembre|januari|februari|maart|april|kunnen|juni|juli|augustus|september|oktober|november|december)(?=\s+(202[1-9]))'
    replaced_title = re.sub(pattern, r'\1&#160;\3', title, flags=re.IGNORECASE)
    return replaced_title

# ✅ Load Excel ref IDs by filename key
def load_refids(excel_file):
    df = pd.read_excel(excel_file)
    refid_dict = {}
    for index, row in df.iterrows():
        filename = str(row['filename']).strip()
        refid_1 = str(row['refid 1']).strip()
        refid_2 = str(row['refid 2']).strip()
        refid_3 = str(row['refid 3']).strip()
        refid_4 = str(row['refid 4']).strip()

        if not filename.lower().endswith(".docx"):
            filename += ".docx"  # Ensure it matches actual filenames

        refid_dict[filename] = {
            'refid_1': refid_1,
            'refid_2': refid_2,
            'refid_3': refid_3,
            'refid_4': refid_4
        }
    return refid_dict

def convert_to_xml(input_file, output_file, file_count, expert_area, refid_1, refid_2, refid_3, refid_4):
    try:
        doc = Document(input_file)
        output_folder = os.path.join(os.path.dirname(input_file), "DummyOut")

        if not os.path.exists(output_folder):
            os.makedirs(output_folder)

        anchor_name_content = doc.paragraphs[0].text.strip() if len(doc.paragraphs) > 0 else "Untitled Anchor"
        title_of_file = doc.paragraphs[1].text.strip() if len(doc.paragraphs) > 1 else "Untitled Title"
        title_of_file = replace_non_breakable_space(title_of_file)

        detected_lang = detect_language_from_paragraphs(doc.paragraphs, 3, 10)
        if detected_lang not in ['nl', 'fr', 'en']:
            detected_lang = ''

        article = ET.Element("article", expert_area=expert_area, xml_lang=refid_3, text_type="", toc="", anchor_name=anchor_name_content)
        title_element = SubElement(article, "title", short=title_of_file)
        title_element.text = title_of_file

        SubElement(article, "citref")
        SubElement(article, "body")

        sourceblock = SubElement(article, "sourceblock")
        SubElement(sourceblock, "source").text = "of/ou"
        SubElement(sourceblock, "source").text = "www.monKEY.be"
        source3 = SubElement(sourceblock, "source")
        url1 = SubElement(source3, "url", attrib={"xlink:href": "http://www.monKEY.be", "refid": refid_1})
        url1.text = "Zoekterm"
        url1.tail = "/"
        url2 = SubElement(source3, "url", attrib={"xlink:href": "http://www.monKEY.be", "refid": refid_2})
        url2.text = "terme de recherche"
        url2.tail = f": FJF {anchor_name_content}"

        tree = ET.ElementTree(article)
        output_file = os.path.join(output_folder, os.path.basename(input_file).replace(".docx", ".xml"))

        with open(output_file, 'wb') as f:
            f.write('<?xml version="1.0" encoding="utf-8"?>\n'.encode('utf-8'))
            f.write('<!DOCTYPE article PUBLIC "-//WKB//DTD SL article comment//EN" "sl_comment.dtd">\n'.encode('utf-8'))
            tree.write(f, encoding='utf-8')

        with open(output_file, 'r', encoding='utf-8') as f:
            xml_content = f.read()

        xml_content = xml_content.replace('><', '>\n<')
        xml_content = xml_content.replace('FJF No. ', 'FJF ')
        xml_content = xml_content.replace('&amp;', '&')
        xml_content = xml_content.replace('expert_area', 'expert-area')
        xml_content = xml_content.replace('xml_lang', 'xml:lang')
        xml_content = xml_content.replace('text_type', 'text-type')
        xml_content = xml_content.replace('anchor_name', 'anchor:name')
        xml_content = xml_content.replace('<citref />', '<citref></citref>')
        xml_content = xml_content.replace('<body />', '<body></body>')
        xml_content = xml_content.replace(' text-type=""', ' text-type="jurisprudence"')
        xml_content = xml_content.replace(' toc=""', ' toc="no"')

        # xml_content = re.sub(r'(FJF \d{4}/\d+)(</source>)', r'"\1"\2', xml_content)
        # xml_content = re.sub(r'(FJF \d{4}/\d+)(</source>)',r'&#8220;\1&#8221;\2',xml_content)
        xml_content = re.sub(r'(FJF \d{4}/\d+)(</source>)',r'“\1”\2',xml_content)


        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(xml_content)

        print(f"Conversion completed for {input_file}")
    except Exception as e:
        messagebox.showerror("Error", f"Error converting {input_file}: {str(e)}")

def select_folder():
    foldername = filedialog.askdirectory()
    entry_folder.delete(0, tk.END)
    entry_folder.insert(0, foldername)

def select_excel_file():
    excel_file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls")])
    entry_excel.delete(0, tk.END)
    entry_excel.insert(0, excel_file_path)

def convert_folder():
    folder = entry_folder.get()
    expert_area = var_expert_area.get()
    excel_file = entry_excel.get()

    if not folder:
        messagebox.showerror("Error", "Please Select a folder Path.")
        return
    if not excel_file:
        messagebox.showerror("Error", "Please Select the Excel file with ref IDs.")
        return

    refid_dict = load_refids(excel_file)

    file_count = 0
    for filename in os.listdir(folder):
        if filename.endswith(".docx"):
            input_file = os.path.join(folder, filename)
            output_folder = os.path.join(os.path.dirname(input_file), "DummyOut")
            output_file = os.path.join(output_folder, filename.replace(".docx", ".xml"))

            refid_info = refid_dict.get(filename)
            if refid_info:
                refid_1 = refid_info['refid_1']
                refid_2 = refid_info['refid_2']
                refid_3 = refid_info['refid_3']
                refid_4 = str(int(float(refid_info['refid_4'])))
            else:
                refid_1 = ""
                refid_2 = ""
                refid_3 = ""
                refid_4 = ""

            print(f"Processing file: {filename}")
            print(f"Refid 1: {refid_1}, Refid 2: {refid_2}, Refid 3: {refid_3}, Refid 4: {refid_4}")

            convert_to_xml(input_file, output_file, file_count, expert_area, refid_1, refid_2, refid_3, refid_4)
            file_count += 1

    messagebox.showinfo("Completed", "Conversion completed for all files.")

# ------------------ GUI ------------------
root = tk.Tk()
root.title("Dummy Creation Tool")
root.geometry("400x500")
root.configure(bg="#f0f0f0")

style = ttk.Style()
style.theme_use("clam")
style.configure('TLabel', background="#f0f0f0", font=("Helvetica", 12))
style.configure('TButton', background="#0078d7", foreground="white", font=("Helvetica", 10), padding=6)
style.configure('TEntry', font=("Helvetica", 10), padding=6)
style.configure('TOptionMenu', font=("Helvetica", 10))

label_folder = ttk.Label(root, text="Select a folder:")
label_folder.pack(pady=10)
entry_folder = ttk.Entry(root, width=50)
entry_folder.pack(pady=10)
browse_button = ttk.Button(root, text="Browse", command=select_folder)
browse_button.pack(pady=10)

label_excel = ttk.Label(root, text="Select Excel file with ref IDs:")
label_excel.pack(pady=10)
entry_excel = ttk.Entry(root, width=50)
entry_excel.pack(pady=10)
browse_excel_button = ttk.Button(root, text="Browse", command=select_excel_file)
browse_excel_button.pack(pady=10)

label_expert_area = ttk.Label(root, text="Select Expert Area:")
label_expert_area.pack(pady=10)
expert_areas = ['FIFI']
var_expert_area = StringVar(root)
var_expert_area.set(expert_areas[0])
dropdown_expert_area = ttk.OptionMenu(root, var_expert_area, *expert_areas)
dropdown_expert_area.pack(pady=10)

convert_button = ttk.Button(root, text="Convert", command=convert_folder)
convert_button.pack(pady=20)

root.mainloop()
