import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
import datetime
import xml.etree.ElementTree as ET
from xml.etree.ElementTree import SubElement
import re
import langdetect
import datefinder
import os
import pandas as pd
from langdetect import detect

current_year = datetime.datetime.now().year

def extract_bold(para):
    return [run.text for run in para.runs if run.bold]

def convert_to_vet_tags(text):
    return f'<vet>{text}</vet>'

def extract_italic(para):
    return [run.text for run in para.runs if run.italic]

def convert_to_it_tags_1(text):
    return f'<it>{text}</it>'

def convert_to_it_tags_2(text):
    return f'<lat>{text}</lat>'

def extract_date(title):
    match = re.search(r'\b\d{1,2}(st|nd|rd|th|er|ème)? \w+ \d{4}\b', title)
    return match.group() if match else None

def detect_language(text):
    try:
        return langdetect.detect(text)
    except:
        return None

def load_refids_by_filename(excel_file):
    df = pd.read_excel(excel_file)
    refid_dict = {}
    df['filename'] = df['filename'].astype(str).str.strip()
    for _, row in df.iterrows():
        filename = row['filename'].strip()
        refid_dict[filename] = {
            'refid_0': f"{int(float(row['refid 0'])):03}",
            'refid_1': str(row['refid 1']).strip(),
            'refid_2': str(row['refid 2']).strip(),
            'refid_3': str(row['refid 3']).strip(),
            'refid_4': str(int(float(row['refid 4']))),
        }
    return refid_dict

# === FULL ORIGINAL convert_to_xml() FUNCTION INSERTED ===

def convert_to_xml(input_file, output_file, refid_1, refid_2, refid_3, refid_4, refid_0):
    try:
        # Load the Word document
        doc = Document(input_file)
        
        output_folder = os.path.join(os.path.dirname(input_file), "output")
        if not os.path.exists(output_folder):
            os.makedirs(output_folder)
    #root Element
        jurisr = ET.Element("jurisr", rsleutel=refid_4)
            
    # create date variable    
        title = doc.paragraphs[1].text
        date = extract_date(title)

        # chrdatum section
        chrdatum = ET.SubElement(jurisr, "chrdatum")
        chardatum = f"{date}"
        chrdatum.text = chardatum

        

    # Store paragraphs processed in the samenst section
        processed_paragraphs_samenst = set()

        # gerecht and samenst section
        gerecht = ET.SubElement(jurisr, "gerecht" , aard="rvst")
        samenst = SubElement(gerecht, "samenst")

        for para in doc.paragraphs[2:]:

            
            
            if para.text.strip():

                if ':' in para.text and 'Noot' not in para.text and 'Note' not in para.text:
                            # Process paragraphs that meet the existing condition
                            parts = para.text.split(':')
                            functie = SubElement(samenst, 'functie')
                            functie.text = parts[0].strip() + ':'
                            al = SubElement(samenst, 'al')

                            italic_text = extract_italic(para)

                            if italic_text:
                                para_text = para.text
                                for text in italic_text:
                                    para_text = para.text.replace(text, convert_to_it_tags_1(text))
                                al.text = para_text
                            
                            
                                # al_text = convert_to_it_tags_1(parts[1].strip()) 
                                # al.text = al_text
                            else:
                                al.text = parts[1].strip()

                            processed_paragraphs_samenst.add(para.text.strip())

                            

                else:
                            # Handle paragraphs that don't meet the existing condition
                            casco_pattern = re.search(r'\b[A-Z][\w\-\'’\(\)\s]*\b(?:\s–\s[A-Z][\w\-\'’\(\)\s]*\b)+', para.text)
                            if casco_pattern:
                                # If the pattern is found, break out of the loop
                                break
                            else:
                                # Process the paragraph
                                al = SubElement(samenst, 'al')

                                italic_text = extract_italic(para)
                                if italic_text:
                                    para_text = para.text
                                    for text in italic_text:
                                        para_text = para.text.replace(text, convert_to_it_tags_1(text))
                                    al.text = para_text
                                else:

                                    al.text = para.text.strip()
                                processed_paragraphs_samenst.add(para.text.strip())    


        lang_jurisr = detect_language(para.text)

        if refid_3 == lang_jurisr:
            jurisr.set('xml_lang', lang_jurisr)
        else:
            jurisr.set('xml_lang', refid_3)        
        # lang_jurisr = detect_language(para.text)
        # if lang_jurisr:
        #     jurisr.set('xml_lang', lang_jurisr)            


        # Store paragraphs processed in the cascobl section
        

            
        # Maintain a set of processed paragraphs for <samenv> and <cascobl>
        processed_paragraphs_samenv = set()  

        processed_paragraphs_casco = set()

        # cascobl section
        if lang_jurisr == "nl":
            cascobl_nl = ET.SubElement(jurisr, "cascobl", xml_lang="nl")
            samenv_nl = ET.SubElement(jurisr, "samenv", xml_lang="nl")
            cascobl_fr = ET.SubElement(jurisr, "cascobl", xml_lang="fr")
            samenv_fr = ET.SubElement(jurisr, "samenv", xml_lang="fr")

        elif lang_jurisr == "fr":
            cascobl_fr = ET.SubElement(jurisr, "cascobl", xml_lang="fr")
            samenv_fr = ET.SubElement(jurisr, "samenv", xml_lang="fr")
            cascobl_nl = ET.SubElement(jurisr, "cascobl", xml_lang="nl")
            samenv_nl = ET.SubElement(jurisr, "samenv", xml_lang="nl")
        else:
            cascobl_nl = ET.SubElement(jurisr, "cascobl", xml_lang="nl")
            samenv_nl = ET.SubElement(jurisr, "samenv", xml_lang="nl")
            cascobl_fr = ET.SubElement(jurisr, "cascobl", xml_lang="fr")
            samenv_fr = ET.SubElement(jurisr, "samenv", xml_lang="fr")

        

        jurtekst_nl = SubElement(samenv_nl, 'jurtekst')
        jurtekst_fr = SubElement(samenv_fr, 'jurtekst')

        for para in doc.paragraphs:
            casco_pattern = re.search(r'\b[A-Z][\w\-\'’\(\)\s]*\b(?:\s–\s[A-Z][\w\-\'’\(\)\s]*\b)+', para.text)
            if casco_pattern and title not in para.text:
                # Determine language of casco text
                lang = detect_language(para.text)
                if lang == "nl":
                    casco = SubElement(cascobl_nl, 'casco')
                    cascobl_nl.set('xml_lang', 'nl')
                elif lang == "fr":
                    casco = SubElement(cascobl_fr, 'casco')
                    cascobl_fr.set('xml_lang', 'fr')
                else:
                    # Default to NL if language detection fails or language is not NL or FR
                    casco = SubElement(cascobl_nl, 'casco')
                    cascobl_nl.set('lang', 'nl')

                # Process the paragraph and assign it to the appropriate casco element
                italic_text = extract_italic(para)
                if italic_text:
                    processed_text = para.text
                    for text in italic_text:
                        processed_text = processed_text.replace(text, convert_to_it_tags_2(text))
                    casco.text = processed_text
                else:
                    casco.text = para.text.strip()

        
                processed_paragraphs_casco.add(para.text.strip())

        # Iterate through paragraphs to identify and separate NL and FR samenv elements
        # Iterate through paragraphs
        for para in doc.paragraphs:
        # Extract anchor name to identify relevant paragraphs for samenv
            # anchor_match = re.search(r'No\.\s\d{4}/\d{1,2}', para.text)
            # if anchor_match is None and title not in para.text:
            if title not in para.text:
                if para.text.strip() not in processed_paragraphs_casco and para.text.strip() not in processed_paragraphs_samenst:
                    # Determine language of samenv text
                    lang_samenv = detect_language(para.text)
                    if lang_samenv == "fr":
                        jurtekst = jurtekst_fr
                    elif lang_samenv == "nl":
                        jurtekst = jurtekst_nl
                    else:
                        # Default to NL if language detection fails or language is not NL or FR
                        if lang_jurisr == "fr":
                            jurtekst = jurtekst_nl or jurtekst_fr
                        else:
                            jurtekst = jurtekst_fr or jurtekst_nl
                        # jurtekst = jurtekst_fr or jurtekst_nl
                        # continue
                        

                    # Split the text into paragraphs using regular expression to handle different line breaks
                    paragraphs = re.split(r'\r?\n\r?\n', para.text.strip())

                    # Create appropriate elements and append to the relevant samenv section
                    for paragraph in paragraphs:
                        # Determine the number of paragraph breaks in the paragraph
                        paragraph_breaks = paragraph.count('\n')

                        # Create appropriate elements based on the number of paragraph breaks
                        if paragraph_breaks > 1:
                            # If there are one or more paragraph breaks, start a new <p> element
                            p = SubElement(jurtekst, 'p')
                            processed_lines = paragraph.split('\n')
                            for line in processed_lines:
                                al = SubElement(p, 'al')
                                al.text = line.strip()
                        else:
                            # If there is no paragraph break, create a new <al> element
                            al = SubElement(jurtekst, 'al')
                            al.text = paragraph.strip()

                        # Process the paragraph and assign it to the appropriate element
                        processed_text = paragraph.strip()
                        italic_text = extract_italic(para)
                        bold_text = extract_bold(para)
                        # superscript_text = extract_super(para)

                        # Replace italic text with <it> tags
                        for text in italic_text:
                            processed_text = processed_text.replace(text, convert_to_it_tags_1(text))

                        # Replace bold text with <vet> tags
                        for text in bold_text:
                            processed_text = processed_text.replace(text, convert_to_vet_tags(text))

                        # for text in superscript_text:
                        #     processed_text = processed_text.replace(text, convert_to_super_tags(text))

                        # Assign the processed text to the appropriate element
                        al.text = processed_text.strip()

                        # Add the processed paragraph to the set of processed paragraphs for the Samenv section
                        processed_paragraphs_samenv.add(para.text.strip())   

# Uitispark Section  

        uitspraak = ET.SubElement(jurisr, "uitspraak")
        jurtekst = SubElement(uitspraak, "jurtekst")
        p = SubElement(jurtekst, "p")
        tabblok = SubElement(p, "tabblok")
        table = SubElement(tabblok, "table", frame="none")
        tgroup = SubElement(table, "tgroup", cols="2")
        colspec1 = SubElement(tgroup, "colspec", align="left", colname="col1", colnum="1", colsep="0", colwidth="5*", rowsep="0")
        colspec2 = SubElement(tgroup, "colspec", align="left", colname="col2", colnum="2", colsep="0", colwidth="5*", rowsep="0")
        tbody = SubElement(tgroup, "tbody", valign="top")
        row = SubElement(tbody, "row")
        if lang_jurisr == "fr":
            entry1 = SubElement(row, "entry")
            al1 = SubElement(entry1, "al")
            al1.text = "Vous pouvez consulter le texte intégral via"
            al2 = SubElement(entry1, "al")
            figblok1 = SubElement(al2, "figblok")
            figuur1 = SubElement(figblok1, "figuur", bestand="QR-CODE", figure_name=f"{current_year}_{refid_0}_FR.gif")
            entry2 = SubElement(row, "entry")
            al3 = SubElement(entry2, "al")
            al3.text = "U kan de integrale tekst raadplegen via"
            al4 = SubElement(entry2, "al")
            figblok2 = SubElement(al4, "figblok")
            figuur2 = SubElement(figblok2, "figuur", bestand="QR-CODE", figure_name=f"{current_year}_{refid_0}_NL.gif")

        elif lang_jurisr == "nl":
                
            entry2 = SubElement(row, "entry")
            al3 = SubElement(entry2, "al")
            al3.text = "U kan de integrale tekst raadplegen via"
            al4 = SubElement(entry2, "al")
            figblok2 = SubElement(al4, "figblok")
            figuur2 = SubElement(figblok2, "figuur", bestand="QR-CODE", figure_name=f"{current_year}_{refid_0}_NL.gif")

            entry1 = SubElement(row, "entry")
            al1 = SubElement(entry1, "al")
            al1.text = "Vous pouvez consulter le texte intégral via"
            al2 = SubElement(entry1, "al")
            figblok1 = SubElement(al2, "figblok")
            figuur1 = SubElement(figblok1, "figuur", bestand="QR-CODE", figure_name=f"{current_year}_{refid_0}_FR.gif")

        elif lang_jurisr == "af":
            entry1 = SubElement(row, "entry")
            al1 = SubElement(entry1, "al")
            al1.text = "Vous pouvez consulter le texte intégral via"
            al2 = SubElement(entry1, "al")
            figblok1 = SubElement(al2, "figblok")
            figuur1 = SubElement(figblok1, "figuur", bestand="QR-CODE", figure_name=f"{current_year}_FR.gif")
            entry2 = SubElement(row, "entry")
            al3 = SubElement(entry2, "al")
            al3.text = "U kan de integrale tekst raadplegen via"
            al4 = SubElement(entry2, "al")
            figblok2 = SubElement(al4, "figblok")
            figuur2 = SubElement(figblok2, "figuur", bestand="QR-CODE", figure_name=f"{current_year}_NL.gif")


                

        # Write XML to file
        
        tree = ET.ElementTree(jurisr)
        tree.write(output_file, encoding='utf-8', xml_declaration=True)
        print(f"Conversion completed for {input_file}")
        
    

    # Write XML to file with line breaks
        xml_content = ET.tostring(jurisr, encoding='utf-8')
        xml_content = xml_content.decode('utf-8')
        xml_content = xml_content.replace('<jurisr', '<?xml version="1.0" encoding="utf-8"?>\n<jurisr')
        
        xml_content = xml_content.replace('<p><al /></p>', '')
    #samenst replacement
        xml_content = xml_content.replace('</functie><al>', '</functie><al>&#x00A0;')

        
    #All tags breaking
        xml_content = xml_content.replace('><', '>\n<')

        

    #replace all the tags with <it> and </it>    
        xml_content = xml_content.replace("&lt;it&gt;", "<it>")
        xml_content = xml_content.replace("&lt;/it&gt;", "</it>")
        xml_content = xml_content.replace("&lt;lat&gt;", "<lat>")
        xml_content = xml_content.replace("&lt;/lat&gt;", "</lat>")
    #replace all the tags with <vet> and </vet>    
        xml_content = xml_content.replace("&lt;vet&gt;", "<vet>")
        xml_content = xml_content.replace("&lt;/vet&gt;", "</vet>")
    #replace <al />
        xml_content = xml_content.replace("<jurtekst>", "<jurtekst>\n<p>")
        xml_content = xml_content.replace("</jurtekst>", "</p>\n</jurtekst>")
        xml_content = xml_content.replace("<al />\n<al>", "</p>\n<p><al>")
        xml_content = xml_content.replace("<al />\n", "")
        xml_content = xml_content.replace(">\n<p>\n</p>\n<", ">\n<")
        xml_content = xml_content.replace("<p>\n<p>", "<p>")
        xml_content = xml_content.replace("</p>\n</p>", "</p>")

        xml_content = xml_content.replace("<it><it>", "<it>")
        xml_content = xml_content.replace("</it></it>", "</it>")
        xml_content = xml_content.replace("<vet> </vet>", " ")
        xml_content = xml_content.replace("<it> </it>", " ")
        xml_content = xml_content.replace("</it><it>", "")
        xml_content = xml_content.replace("<it>.</it>", ".")

        xml_content = xml_content.replace("1er", "1<super>er</super>")
        xml_content = xml_content.replace("1re", "1<super>re</super>")

        xml_content = xml_content.replace("figure_name", "figure-name")

        xml_content = xml_content.replace('xml_lang', 'xml:lang')

    #Replacements for only Fr files
        if lang_jurisr == "fr":
            xml_content = xml_content.replace("«", "«&#160;")

    # Regex pattern to match the date pattern and put &#160; between date and month
        d_pattern = re.compile(r'(\d{1,2})(\s+)?(janvier|février|mars|avril|peut|juni|juin|juillet|août|septembre|octobre|novembre|décembre|januari|februari|maart|april|kunnen|juni|juli|augustus|september|oktober|november|december)(?=\s+(202[1-9]))')

        # Function to perform the replacement
        def replace_date(match):
            day = match.group(1)
            month = match.group(3)
            return f'{day}&#160;{month}'

        # Perform the replacement
        xml_content = re.sub(d_pattern, replace_date, xml_content)

        # xml_content = xml_content.replace(r'<jurtekst>\n<p>\n<al>No\. \d{4}/\d{2}</al>\n</p>\n<p>', "<jurtekst>\n<p>")

    # Regex pattern to match the specific segment
        # pattern = re.compile(r'<jurtekst>\n<p>\n<al>No\. \d{4}/\d{2}</al>\n</p>\n<p>')
    # For remove unnessasary anchor names in samenst section
        pattern = re.compile(r'<jurtekst>\n<p>\n<al>No. 20([1-9][0-9])/(.*?)</al>\n</p>\n<p>')
    # Replace the matched segment with the desired string
        replacement = '<jurtekst>\n<p>'
        xml_content = re.sub(pattern, replacement, xml_content)



    #empty <it> tags replacement
        # xml_content = xml_content.replace("r'<it>\s*</it>'", " ")
    
        
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(xml_content)
        print(f"Conversion completed for {input_file}")
    except Exception as e:
        messagebox.showerror("Error", f"Error converting {input_file}: {str(e)}")

# === END OF convert_to_xml() FUNCTION ===

# Folder and File Selection UI
def select_folder():
    foldername = filedialog.askdirectory()
    entry_folder.delete(0, tk.END)
    entry_folder.insert(0, foldername)

def select_excel_file():
    excel_file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls")])
    entry_excel.delete(0, tk.END)
    entry_excel.insert(0, excel_file_path)

def convert_folder():
    folder = entry_folder.get()
    excel_file = entry_excel.get()

    if not folder:
        messagebox.showerror("Error", "Please select a folder.")
        return
    if not excel_file:
        messagebox.showerror("Error", "Please select the Excel file with ref IDs.")
        return

    refid_dict = load_refids_by_filename(excel_file)

    for filename in sorted(os.listdir(folder)):
        if filename.endswith(".docx") and not filename.startswith("~$"):
            input_file = os.path.join(folder, filename)
            output_folder = os.path.join(folder, "output")
            output_file = os.path.join(output_folder, filename.replace(".docx", ".xml"))

            refid_info = refid_dict.get(filename)
            if not refid_info:
                print(f"Warning: No refid data found for {filename}")
                continue

            refid_1 = refid_info['refid_1']
            refid_2 = refid_info['refid_2']
            refid_3 = refid_info['refid_3']
            refid_4 = refid_info['refid_4']
            refid_0 = refid_info['refid_0']

            print(f"Processing file: {filename}")
            print(f"Refid 1: {refid_1}, Refid 2: {refid_2}, Refid 3: {refid_3}, Refid 4: {refid_4}, Refid 0: {refid_0}")

            convert_to_xml(input_file, output_file, refid_1, refid_2, refid_3, refid_4, refid_0)

    messagebox.showinfo("Completed", "Conversion completed for all files.")

# Tkinter UI Setup
root = tk.Tk()
root.title("Innodata FJF Automation Tool")
root.geometry("500x400")
root.configure(bg="#f0f4ff")

header = tk.Label(root, text="Innodata FJF Automation Tool", font=("Helvetica", 16, "bold"), fg="#0056b3", bg="#f0f4ff")
header.pack(pady=10)

label_folder = tk.Label(root, text="Select a folder:", font=("Arial", 12), bg="#f0f4ff", fg="#0056b3")
label_folder.pack(pady=5)
entry_folder = tk.Entry(root, width=50, bd=2)
entry_folder.pack(pady=5)
browse_button = tk.Button(root, text="Browse", command=select_folder, bg="#0056b3", fg="white", bd=0, padx=10, pady=5)
browse_button.pack(pady=5)

label_excel = tk.Label(root, text="Select Excel file with ref IDs:", font=("Arial", 12), bg="#f0f4ff", fg="#0056b3")
label_excel.pack(pady=5)
entry_excel = tk.Entry(root, width=50, bd=2)
entry_excel.pack(pady=5)
browse_excel_button = tk.Button(root, text="Browse", command=select_excel_file, bg="#0056b3", fg="white", bd=0, padx=10, pady=5)
browse_excel_button.pack(pady=5)

convert_button = tk.Button(root, text="Convert", command=convert_folder, bg="#0056b3", fg="white", font=("Arial", 12, "bold"), bd=0, padx=20, pady=10)
convert_button.pack(pady=10)

root.mainloop()
